import os
import tempfile
import json
from langchain_core.documents import Document
import docx
import zipfile
import lancedb
import pyarrow as pa
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter

from web.documents.utils.custom_embeddings import CustomEmbeddings

DB_PATH = './web/documents/lancedb_storage'
TABLE_NAME = 'my_knowledge_base'


def load_docx(path):
    try:
        with zipfile.ZipFile(path):
            pass
    except zipfile.BadZipFile:
        raise ValueError("文件不是有效的 .docx 格式，请上传 Word 2007+ 的 .docx 文件")
    doc = docx.Document(path)
    text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=text, metadata={"source": path})]


def insert_documents(file, knowledgeId,characterId):
    print("insert")
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.name)[-1]) as tmp:
        for chunk in file.chunks():
            tmp.write(chunk)
        tmp_path = tmp.name

    ext = os.path.splitext(file.name)[-1].lower()
    if ext == ".txt":
        loader = TextLoader(tmp_path, encoding='utf-8')
        documents = loader.load()
    elif ext == ".pdf":
        loader = PyPDFLoader(tmp_path)
        documents = loader.load()
    elif ext == ".docx":
        documents = load_docx(tmp_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")

    os.unlink(tmp_path)

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    texts = text_splitter.split_documents(documents)

    print("切分：", len(texts))

    # 生成向量
    embeddings = CustomEmbeddings()
    contents = [t.page_content for t in texts]
    vectors = embeddings.embed_documents(contents)

    # 直接用 pyarrow 写入，knowledgeId 作为独立列
    data = pa.table({
        "vector": pa.array(vectors, type=pa.list_(pa.float32())),
        "text": pa.array(contents, type=pa.string()),
        "knowledgeId": pa.array([str(knowledgeId)] * len(texts), type=pa.string()),
        "characterId": pa.array([str(characterId)] * len(texts), type=pa.string()),
    })

    db = lancedb.connect(DB_PATH)
    if TABLE_NAME in db.table_names():
        table = db.open_table(TABLE_NAME)
        table.add(data)
    else:
        db.create_table(TABLE_NAME, data=data)

    table = db.open_table(TABLE_NAME)
    # 输出刚插入的内容
    df = table.to_pandas()
    recent = df.tail(len(texts))
    print(recent[['text', 'knowledgeId', 'characterId']])

    print("insert :", table.count_rows())


def delete_documents(knowledgeId):
    db = lancedb.connect(DB_PATH)

    if TABLE_NAME not in db.table_names():
        print("[delete] 表不存在，无需删除")
        return

    table = db.open_table(TABLE_NAME)
    before = table.count_rows()
    table.delete(f"knowledgeId = '{knowledgeId}'")  # ✅ 现在是独立列，可以直接删
    after = table.count_rows()

    print(f"[delete] knowledgeId={knowledgeId} 删除了 {before - after} 条，剩余 {after} 条")